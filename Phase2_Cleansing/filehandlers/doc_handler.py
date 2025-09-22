''' This function opens a .docx file, scans each paragraph for PII, cleans the text,
replaces it in the document, logs what was found,
and saves a sanitized version of the file.
'''



import docx

from ..detectors import detect_pii_in_text
from ..maskers import mask_text
from ..audit  import AuditLogger
audit = AuditLogger("audit_log.csv")

#from ..audit import write_audit_row

def clean_doc_file(input_path, output_path, action, use_spacy, audit):
    try:
        doc = docx.Document(input_path)  # loads the word document from path to doc
    except FileNotFoundError:
        print(f"[FAIL] Word file not found: {input_path}")
        return False
    except Exception as e:
        print(f"[FAIL] Could not open DOCX file {input_path}: {e}")
        return False

    try:
        for p_idx, para in enumerate(doc.paragraphs): # goes through  each para in the doc
            if not para.text or not para.text.strip(): # skip the para, if it has no text or empty
                continue
            detections = detect_pii_in_text(para.text, use_spacy=use_spacy)  # scan para for PII
            if detections: # if found, mask it and store the cleaned version in cleaned_text
                cleaned_text = mask_text(para.text, detections, action=action)
                for run in para.runs: # a paragraph may consist of multiple runs (chunks of text with different formatting: bold, italic, etc.
                    run.text = ""   # Clear all existing runs
                para.add_run(cleaned_text)  # Add back a single run containing the cleaned text
                for d in detections:  # log each detection
                    audit.write_row(input_path, output_path,
                                    d.get("source"), d.get("type"), d.get("match"),
                                    action, notes = f"paragraph:{p_idx+1}")

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"[FAIL] Error processing DOCX file {input_path}: {e}")
        return False